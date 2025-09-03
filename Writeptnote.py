'''

Part of CHRIS that generates notes.

Ontology of self.note parts:

Titles: CHRIS will format all simple (not tables or dicts) values as titles if
and only if the key contains 'title' (case insensitive). In text notes, titlesl
will be formatted with a following line filled with '=' resulting in a double
underlined title.

Tables: CHRIS will format all dictionary values as tables if and only if the
key contains 'table' (case insensitive). Only simple values are allowed in
tables. No titles, dictionaries, or other tables within a table. Keys will be
inserted into the first column and their corresponding values in the second. In
text notes, the result will be lines with the key on the left then a colon and
the value on the right.

Paragraphs: If the key of a value does not contain 'title' or 'table' (case
insensitive), and the value is not a dict, then the value will be interpreted
as a paragraph of text. If there is a newline in the value, then a paragraph
break (not a simple line break) will be inserted there.

Dictionaries: CHRIS will consider all dictionary values as containers. No
formatting is applied when CHRIS encounters a dictionary, but it will
recursively read all the keys and values from the dictionary.

'''

import re
import matplotlib #will be used to build graphs of surveys
from itertools import cycle
import yaml
import logging
import sys
import traceback


from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

from debugging import format_stacktrace

# These variables should not be used anymore. They are not read in from the
# surveys and should not be asked for by the note.
svar_blacklist =[
    "moveInvolMovSev",
    "moveHyg",
    "moveHobby",
    "moveTurn",
    "moveGetUp",
    "moveFreeze",
    "socMuch",
    "dietType",
    "moveWalkBal",
]


##To-do: add in functionality if the patient fills out the form themselves.
#honestly, could just write in replacing careName with the ptName into the carename function?

#add in a probSpeed section in bigCHRIS, as no note variable exists with this info yet.

#Actually, need to put ALL variables in 'newvariablelist' of the ptsurvey class in.



class Ptnote:

    domainCC = {
        'npiq' : "At present, the history provided suggests that the dominant problem may be behavioral or neuropsychiatric.  ",
        'ecogVisuosp' : "At present, the history provided suggests that the dominant problem may be visuospatial.  ",
        'ecogLang' : "At present, the history provided suggests that the dominant problem may involve language.  ",
        'ecogMem' : "At present, the history provided suggests that the dominant problem may involve memory.  ",
        'ecogExec' : "At present, the history provided suggests that the dominant problem may be executive function.  ",
        }

        # Read the dictionary from survey.yaml
    with open('ptnote.yaml') as f:
        standardizedsurveys = yaml.safe_load(f)

    def __init__(self, survey):

        #assigns data from the survey class instance to the variable "survey"
        self.survey = survey
        #creates empty self.note and self.variables variables, so the functions below won't bug out
        self.note = {}
        self.variables = {}
        self.whole_note = ""

        #this section renames functions so that you don't have to type self.x() each time, you can just type x()
        findvar = self.findvar
        # ~ findpath = self.findpath
        write = self.write
        nvar = self.nvar
        svar = self.svar
        contains_response = self.contains_response
        anysurveyresponses = self.anysurveyresponses
        buildlist = self.buildlist
        buildsurveydesc = self.buildsurveydesc
        domain = self.domain
        score = self.score
        score0to1 = self.score0to1
        surveyscore = self.surveyscore
        surveypercent = self.surveypercent
        compositescore = self.compositescore
        compositepercent = self.compositepercent
        genderpronouns = self.genderpronouns
        an = self.an
        table = self.table
        # ~ export = self.export
        iscomplete = self.iscomplete
        allvariables = self.allvariables
        replacenvars = self.replacenvars
        replacebuildlists = self.replacebuildlists

        # You are not allowed to use nvar() in this data structure!
        self.variables = {

            'hpi' : {

                'introvariables' : {

                    'ptHand' : 'ambidextrous' if "bit of both" in svar('ptHand') else svar('ptHand'),
                    'ptHandDesc' : " (" + svar('ptHandDesc') + ")" if "bit of both" in svar('ptHand') else "",
                    'probSum' : "no significant past medical history " if "none of the above" in svar('probList') else "a past medical history of ",
                    'probList' : str(svar('probList')),
                    'careFreq' : svar('careFreq').lower(),
                    'careHr' : svar('careHr').lower(),
                    'careSum' : svar('careName') + ", " + svar('ptId') + "'s " + svar('careRelate').lower() + ", " + "helped to provide the history.  " + svar('careName') + " sees " + svar('ptId') + " about " + svar('careFreq') + ", or approximately " + svar('careHr') + ", is " + svar('careAge') + " years old, and has known the patient for " + svar('careYr') + " years. " if score('ptfillYN') == 1 else
                        "",
                    'probSpeed' : 'When asked about how rapidly symptoms have changed, careName reports that they have ' + svar('probSpeed').lower() + ".  " if not contains_response(['probSpeed'], ["There has been no worsening"]) == True else
                        'When asked about how rapidly symptoms have changed, careName reports that they have not changed.  ',
                    'rapidWorsen' : 'not ' if svar('rapidWorsen') == "No" else " "
                },
                'MemDescvariables' : {

                    'ecogMemList' : buildsurveydesc('ecogMem') if buildsurveydesc('ecogMem') != "" else
                        "He/She reported better or no change to all memory Ecog questions.  "  ,
                    'memStart'  : "Memory problems started in approximately " + svar('memStart') + ". " if svar('memStart') != "#memStart#" else
                         "",
                    'memLimit' : svar('memLimit') + "  " if svar('memLimit') != "No memory problems are present." else
                        ""},
                'execDescVariables' : {
                    'ecogExecList' : buildsurveydesc('ecogExec') if buildsurveydesc('ecogExec') != "" else
                        "He/She reported better or no change to all executive function Ecog questions.  "  ,
                    'execJudge' : "A change in judgment is also described." if svar('execJudge') == "Yes, a change" else
                        "",
                    'execStart' : "Executive problems started in approximately " + svar('execStart') + ". " if svar('execStart') != "#execStart#" else
                        "",
                    'execLimit' : svar('execLimit') + "  " if svar('execLimit') != "None of these problems are present." else
                        ""},

                'langDescvariables' : {

                    'ecogLangList' : buildsurveydesc('ecogLang') if buildsurveydesc('ecogLang') != "" else
                        "He/She reported better or no change to all language Ecog questions.  ",
                    'langQual' :
                        str(svar('langQualSevere')) + " Speech quality is described as more: "+ str(svar('langQualChange')) + ". " if svar('langQualChange') != '[]' and not contains_response(['langQualChange'], ["unchanged"]) == True and svar('langQualSevere') != "Not at all (no problems)." else
                        str(svar('langQualSevere')) + " No more specific description of the change is provided. " if (svar('langQualChange') == '[]' or contains_response(['langQualChange'], ['unchanged']) == True) and svar('langQualSevere') != "Not at all (no problems)." else
                        "While speech changes are not described as at all impacting communication, speech quality is described as more: " + svar('langQualChange') + ". " if svar('langQualChange') != '[]' and contains_response(['langQualChange'], ['unchanged']) and svar('langQualSevere') == "Not at all (no problems)." else
                        "No change in speech quality is described. " if (svar('langQualChange') == '[]' or contains_response(['langQualChange'], ['unchanged', "Unchanged"]) == True) and svar('langQualSevere') == "Not at all (no problems)." else
                        "***Speech Quality and Problems***",
                    'langElseList' : {
                    'writeChange' : "writing" if svar('langWrite') == "Yes" else
                        "",
                    'readChange' : "reading" if svar('langRead') == "Yes" else
                        "",
                    'spellChange' : "spelling" if svar('langSpell') == "Yes" else
                        "",
                    'grammarChange' : "grammar" if svar('langGrammar') == "Yes" else
                        "",
                    'spAmountChange' :"the amount of speech he/she produces" if svar('langAmount') == "More" else
                        ""},
                    'langLimit' : svar('langLimit') + "  " if svar('langLimit') != "No such language problems are present" else
                        "",
                    'langStart' : "Language problems started in approximately " + svar('langStart') + ". " if svar('langStart') != "#langStart#" else
                        ""},

                'visuospVariables' : {

                    'ecogVisuospList' : buildsurveydesc('ecogVisuosp') if buildsurveydesc('ecogVisuosp') != "" else
                        "He/She reported better or no change to all visuospatial Ecog questions.  ",
                    'visuospFamPeopleObj' : "Some new problems recognizing familiar faces or objects are also described." if svar('visuospFamPeopleObj') == "Yes" else
                        "No new problems recognizing familiar faces or objects are described.",
                    'visuospStart' : "Visuospatial problems started in approximately " + svar('visuospStart') + ". " if svar('visuospStart') != "#visuospStart#" else
                        ""},

                'moveDescVariables' : {

                    'fallYN' : "" if svar('moveFall') == "Yes" else
                        "not ",
                    'fallCaut' : " " if svar('moveFallCaution') == "Yes" else
                        "not ",
                    'weak' : "some " if svar('moveWeak') == "Yes" else
                        "no ",
                    'fineMot' : "some " if svar('moveFine') == "Yes" else
                        "no ",
                    'moveInvolMovSev' : "These movements do not at all intervere with daily activities.  " if svar('moveInvolMovSev') == "Not at all" else
                        "",
                    'swallow' : "When asked about swallowing, the response is: '" + svar('moveSwallow') + "'  " if svar('moveSwallow') != "There are no problems." else
                        "There are no significant changes in swallowing.  ",
                    'hygMov' : "some " if svar('moveHyg') != "Not at all (no problems)." and svar('moveHyg') != "" else
                        "no ",
                    'hobMov' : "some " if svar('moveHobby') != "Not at all (no problems)." and svar('moveHobby') != "" else
                        "no ",
                    'turnMov' : "some " if svar('moveTurn') != "Not at all (no problems)." and svar('moveTurn') != "" else
                        "no ",
                    'getUpMov' : "some " if svar('moveGetUp') != "Not at all (no problems)." and svar('moveGetUp') != "" else
                        "no ",
                    'freezeMov' : "some " if svar('moveFreeze') != "Not at all (no problems)." and svar('moveFreeze') != "" else
                        "no ",
                    'involMot1' : "not " if "no" or "maybe" in svar('moveInvolMov') else
                        "",
                    'involMot2': "" if  "no" or "maybe" in svar('moveInvolMov') else
                        ", which are difficult to describe. " if contains_response('moveInvolMov', "other") else
                        ", including" + svar('moveInvolMov'),
                    'slownessInit' : "There has been no slowness when initiating movements.  " if svar('slownessInit') == "No" else
                        "There has been slowness when initiating movements.  ",
                    'rigidityMotion' : "There has been rigidity on passive range of motion.  " if svar('rigidityMotion') == "Yes" else
                        'There has been no rigidity on passive range of motion.  ',
                    'lossPost' : "There has been a loss of postural stability.  " if svar('lossPostStability') =="Yes" else
                        "There has been no loss of postural stability.  ",
                    'restTremor' : "There has been a resting tremor.  " if svar('restTremor') == "Yes" else
                        "There has been no resting tremor.  "},

                'senseVariables' : {

                    'numbSens' : "denies any" if svar('senseNumb') != "Yes" else
                        "endorses some",
                    'hear' : " some " if svar('senseHearYN') != "Not at all (no problems)." else
                        "no ",
                    'sight1' : "no " if contains_response(['senseSight'], ["no", "other"]) == True or svar('senseSight') == "[]" else
                        "some ",
                    'sight2' :  "." if contains_response(['senseSight'], ["no", "maybe"]) == True else
                        ", described as follows: " + str(svar('senseSight')) + ".",
                    'smell' : "There is no change in sense of taste or smell. " if contains_response(['senseSmell'], ['no']) == True else
                        "He/she is uncertain about a change in sense of taste or smell. " if contains_response(['senseSmell'], ["maybe", "Other"]) == True else
                        "He/she also reports:" + str(svar('senseSmell')) + ". "},

                'autoVariables' : {

                    'bowel' : "some" if svar('autoBowelIncont') == "Yes" else
                        "no"},

                'behVariables' : {

                    'npiqList' : buildsurveydesc('npiq'),
                    'behInapprop' : 'some ' if svar('behInapprop') == "Maybe" else 'no ' if svar('behInapprop') == "No" else
                    "",
                    'behOralFix' : 'some ' if svar('behOralFix') == "Maybe" else 'no ' if svar('behOralFix') == "No" else
                    "",
                    'behReligChange' : 'some ' if svar('behReligChange') == "Maybe" else 'no ' if svar('behReligChange') == "No" else
                    "",
                    'behIntimacyChange' : 'some ' if svar('behIntimacyChange') == "Maybe" else 'no ' if svar('behIntimacyChange') == "No" else
                    "",
                    'behHygieneChange' : 'some ' if svar('behHygieneChange') == "Maybe" else 'no ' if svar('behHygieneChange') == "No" else
                    "",
                    'labile' : "" if svar('behEmoLabile') != "No" else
                        "not ",
                    'passSI' : "does not deny " if svar('behPassiveSI') != "No" else
                        "denies ",
                    'actSI1' : "does not deny " if svar('behActSI') != "No" else
                        "denies ",
                    'compulse' : "Regarding compulsive behaviors, " + svar('behCompulse').lower() + " ",
                    'actSI2' : "In addition," if svar('behActSI') != "No" and svar('behPassiveSI') != "No" else
                        "However," if svar('behActSI') == "No" and svar('behPassiveSI') != "No" else
                        "Furthermore,",
                    'violent' : "is reportedly " if svar('behViolent') != "No" else
                        "also " if svar('behActSI') != "No" and svar('behViolent') != "No" else
                        "is not reportedly ",
                    'behLimit' : svar('behLimit') + "  " if svar('behLimit') != "No such personality, mood or behavior problems are present" else
                        "",
                    'behStart' : "Behavior problems started in approximately " + svar('behStart') + ". " if svar('behStart') != "#behStart#" else "",
                    #the following six are super redundant!  Optimize me!
                    'readEmoEyes' : "is never " if score('readEmoEyes') == 1 else "generally is not " if score('readEmoEyes') == 2 else "occasionally is not " if score('readEmoEyes') == 3 else "is occasionally " if score('readEmoEyes') == 4 else "is generally " if score('readEmoEyes') == 5 else "is always " if score('readEmoEyes') == 6 else "",
                    'detectFaceExp' : "is never " if score('detectFaceExp') == 1 else "generally is not " if score('detectFaceExp') == 2 else "occasionally is not " if score('detectFaceExp') == 3 else "is occasionally " if score('detectFaceExp') == 4 else "is generally " if score('detectFaceExp') == 5 else "is always " if score('detectFaceExp') == 6 else "",
                    'understandingOthers' : "are never " if score('understandingOthers') == 1 else "generally are not " if score('understandingOthers') == 2 else "occasionally are not " if score('understandingOthers') == 3 else "are occasionally " if score('understandingOthers') == 4 else "are generally " if score('understandingOthers') == 5 else "are always " if score('understandingOthers') == 6 else "",
                    'detectJokeTaste' : "can never " if score('detectJokeTaste') == 1 else "generally can not " if score('detectJokeTaste') == 2 else "occasionally can not " if score('detectJokeTaste') == 3 else "can occasionally " if score('detectJokeTaste') == 4 else "can generally " if score('detectJokeTaste') == 5 else "can always " if score('detectJokeTaste') == 6 else "",
                    'detectInappropriate' : "is never " if score('detectInappropriate') == 1 else "generally is not " if score('detectInappropriate') == 2 else "occasionally is not " if score('detectInappropriate') == 3 else "is occasionally " if score('detectInappropriate') == 4 else "is generally " if score('detectInappropriate') == 5 else "is always " if score('detectInappropriate') == 6 else "",
                    'detectLying' : "is never " if score('detectLying') == 1 else "generally is not " if score('detectLying') == 2 else "occasionally is not " if score('detectLying') == 3 else "is occasionally " if score('detectLying') == 4 else "is generally " if score('detectLying') == 5 else "is always " if score('detectLying') == 6 else ""},

                'sleepVariables' : {

                    'epworthList' : buildsurveydesc('epworth'),
                    'sleepInsomnia' : "There are no problems with insomnia." if svar('sleepInsomnia') == "No problems." else
                        svar('sleepInsomnia'),
                    'sleepEDS' : "No daytime sleepiness was described. " if svar('sleepEDS') == "No daytime sleepiness." else
                        svar('sleepEDS'),
                    'parasom' : "There is some " + str(svar('sleepParasom')) + ".  " if svar('sleepParasom') != "no" and svar('sleepParasom') != "No" and svar('sleepParasom') != "maybe" and svar('sleepParasom') != "Maybe" else
                        "No somnambulism or somniloquy is described.  ",
                    'rsbd' : "Some " if svar('sleepRSBD') == "yes" else
                        "No ",
                    # ~ 'epwFiller' : "When asked in more detail, no further drowsiness was detected on an Epworth scale (Johns Sleep 1991)." if svar('epworthList') == '[]' else
                        # ~ "An Epworth Scale revealed the following",
                    'snore' : "Some snoring " if svar('sleepSnore') == "Yes" else
                        "No snoring ",
                    'osa1' : "but " if svar('sleepOSA') == "Yes" and svar('snore') == "Some snoring" else
                        "and ",
                    'osa2' : "some " if svar('sleepOSA') == "Yes" else
                        "no ",
                    'sleepNap' : svar('sleepNap').lower(),
                    'drowsy' : "but he/she is drowsy or lethargic " + svar('sleepDrowsy') + ". " if contains_response(['sleepDrowsy'], ["Not at all", "not at all"]) == True and (svar('sleepNap') == "Not at all" or svar('sleepNap') == "not at all") else
                        "and he/she is drowsy or lethargic " + svar('sleepDrowsy') + ". " if not contains_response(['sleepDrowsy'], ["Not at all", "not at all"]) == True and (svar('sleepNap') != "Not at all" or svar('sleepNap') != "not at all") else
                        "",
                    'sleepDiffArouse' : svar('sleepDiffArouse').lower()},

                'faqVariables' : {

                    'faqList' : buildsurveydesc('FAQ'),
                    'adlList' : buildsurveydesc('ADL'),
                    'drive' : "currently " if svar('adlDrive') == "Yes" else
                        "not ",
                    'mva' : 'possible ' if svar('adlMVA') == "Maybe" else
                        "" if svar('adlMVA') == "Yes" else
                        "no ",
                    'safeHome' : "possible " if svar('adlHomeSafe') == "Maybe" else
                        "" if svar('adlHomeSafe') == "Yes" else
                        "no "},

                'careDescVariables' : {

                    'zaritList' : buildsurveydesc('zarit'),
                    'careDesc1' : svar('ptId') +"'s " + svar('careRelate').lower() + ", " + svar('careName') + ", answered several questions regarding caregiver burden and stress, starting with the NPI-Q, which suggests the following: " + buildsurveydesc('npiqCaregiver') if buildsurveydesc('npiqCaregiver') != '' else
                    ""}},

            'socHxVariables' : {

                'dM1' : "Regarding developmental milestones, he/she was behind on some milestones.  " if svar('devMilesYN') == "No" else
                    "He/she met all developmental milestones.  ",
                'specEd' : "some " if svar('specEdYN') == "Yes" else
                    "no ",
                'highestEduc' : "less than grade 12" if svar('highestEduc') == "did not graduate high school" else
                    "***" if svar('highestEduc') == "other" else
                    svar('highestEduc'),
                'gradeType' : "***" if svar('gradeType') == "other" else
                    svar('gradeType'),
                'careRelate' : "***" if contains_response(['careRelate'], ["other"]) == True else
                    svar('careRelate'),
                'nearbyFriends' : "" if svar('supportLocal') == "Yes" else
                    "no ",
                'social' : "" if contains_response(['socMuch'], ['yes']) == True else
                    "not ",
                'diet' : "He/she follows a special diet.  " if contains_response(['dietType'], ['yes']) else
                    "He/she follows no particular diet. ",
                'tox' : "He/she believes that he/she has been exposed to environmental hazards.  " if contains_response(['hazard'], ['yes']) else
                    "He/she has not been exposed to any environmental hazards.  ",
                'gun' : "There are guns in the home, reportedly well secured.  " if svar('gunNum') == "two or more" and svar('gunSecure') == "Yes" else
                    "There are guns in the home, not clearly well secured.  " if svar('gunNum') == "two or more" and svar('gunSecure') != "Yes" else
                    "There is a gun in the home, and it is well secured.  " if svar('gunNum') == "one" and svar('gunSecure') == "Yes" else
                    "There is a gun in the home, and it is not clearly well secured.  " if svar('gunNum') == "one" and svar('gunSecure') != "Yes" else
                    "There are no guns in the home.  " if svar('gunNum') == "none" else "",
                'houseMate' : svar('houseMate')},

            'initAssessVariables' : {

              'domCC' : self.domainCC[domain('npiq', 'ecogVisuosp', 'ecogLang', 'ecogMem', 'ecogExec')],

              'dangerList' :

                  {'rapidWorsenDanger' : "rapid and recent decline" if svar('rapidWorsen') != "No" else
                      "",
                  'gunStore' : "guns are not securely stored" if svar('gunSecure') != "Yes" and svar('gunNum') != 'none' else
                      "",
                  'concernViolent' : "some concern for violent behavior towards others" if svar('behViolent') != "No" else
                      "",
                  'concernActSI' : "concern for active suicidal ideation" if svar('behActSI') != "No" else
                      "",
                  'concernPassSI' : "concern for passive suicidal ideation" if svar('behPassiveSI') != "No" else
                      "",
                  # 'safehome' is a note variable and you cannot access those here!
                  # ~ 'concernSafeHome' : "concerns regarding safety at home" if nvar('safeHome') != 'no ' else
                      # ~ "",
                  # moveWalkBal doesn't exist
                  'concernFall' : "some fall risks without clear precautions in place" if (svar('moveFall') != "No" or (svar('moveWalkBal') != "Not at all (no problems).") and svar('moveFallCaution') != "Yes") else
                      "",
                  # moveWalkBal doesn't exist
                  'concernFall2' : "some fall risks though with some precautions in place" if (svar('moveFall') != "No" or (svar('moveWalkBal') != "Not at all (no problems).") and svar('moveFallCaution') == "Yes") else
                      "",
                  'concernSwallow' : "some problems swallowing" if svar('moveSwallow') != "There are no problems." else
                      "",
                  'concernDriving' : "some concern regarding driving" if (svar('adlDrive') == "Yes" and svar('adlMVA') == "Yes") or surveyscore('ecogVisuosp') > 2 or surveyscore('FAQ') > 9 else
                      ""},

              'demRiskList' :

                  {'depRisk' : "depression" if svar('behDepress') != "Not present" or surveyscore('phq') > 5 else
                      "",
                  'anxRisk' : "anxiety" if svar('behAnx') != "Not present" or surveyscore('gad7') > 5 else
                      "",
                  'sleepyRisk' : "potential excess daytime sleepiness" if svar('sleepEDS') != "No daytime sleepiness." or surveyscore('epworth') > 9 else
                      "",
                  "sleepApRisk" : "potential sleep apnea" if svar('sleepOSA') != "No" or svar('sleepSnore') == "Yes" else
                      ""},

              #1-6-20: Note that "The functional activities questionnaire is suggestive of a mild dementia. " was originally if surveyscore('FAQ') > 6, but a secondary condition
              #'and surveyscore('FAQ') <= 15' was added, to further specify and prevent errors.
              'fxnSummary' :  "The functional activities questionnaire is suggestive of severe dementia.  " if surveyscore('FAQ') > 23 else
                  "The functional activities questionnaire is suggestive of at least a moderate dementia.  " if surveyscore('FAQ') > 15 and surveyscore('FAQ') < 23 else
                  "The functional activities questionnaire is suggestive of a mild dementia.  " if surveyscore('FAQ') >= 6 and surveyscore('FAQ') <= 15 else
                   "The functional activities questionnaire does not suggest the presence of a dementia.  ",

              #1-6-20 note: "Basic activities of daily living are impaired. " if surveyscore('ADL') > 3 was changed to "Basic activities of daily living are impaired. " if surveyscore('ADL') >= 3
              #because if ADL = 3, this piece of code would not work.
              'fxnSum2' : "Basic activities of daily living are impaired. " if surveyscore('ADL') >= 3 else
                  "Basic activities of daily living may be impacted. " if surveyscore('ADL') > 1 and surveyscore('ADL') < 3 else
                  "Basic activities of daily living are intact. ",

              'careBurden' : "The response on the Zarit is within the top quartile for caregiver stress.  " if surveyscore('zarit') >= 17 else
                  "The response on the Zarit falls with between the 50th to 75th percentile.  " if surveyscore('zarit') >= 9 and surveyscore('zarit') < 17 else
                  "The response on the Zarit falls with between the 25th to 50th percentile.  " if surveyscore('zarit') >= 3 and surveyscore('zarit') < 9 else
                  "There are no obvious signs of significant caregiver burden.  " if score('ptfillYN') == 1 else
                  "",

              'NPIQCareBurden' : "Neuropsychiatric and behavioral symptoms are significantly contributing to caregiver stress. " if anysurveyresponses('npiqCaregiver', 'npiCareMod', 'npiCareSev', 'npiCareExt') == True else
                  "Neuropsychiatric or behavioral symptoms do not cause significant caregiver burden. " if score('ptfillYN') == 1 else
                  ""
            }
        }

        # Must define variables that use nvar here NOT above.
        self.variables['initAssessVariables']['dangerList']['concernSafeHome'] = "concerns regarding safety at home" if nvar('safeHome') != 'no ' else ""
        self.variables['hpi']['sleepVariables']['epwFiller'] = "When asked in more detail, no further drowsiness was detected on an Epworth scale (Johns Sleep 1991)." if nvar('epworthList') == '[]' else "An Epworth Scale revealed the following: "

        #This dictionary contains the note structure itself.
        self.note = {

            'hpi' : { # history of present illness
                'hpiTitle' : 'History of Present Illness',
                'introoutput' : str(svar('ptId')) + " is a ***" + "-year-old " + str(nvar('ptHand')) + str(nvar('ptHandDesc')) +
                        " " + str(svar('profType').lower()) + " " + "with " + str(nvar('probSum')) + str(svar('probList')) + " who is visiting the Rocky Mountain Memory Disorders Clinic at the University of Colorado, Denver. " +
                        nvar('careSum') + nvar('probSpeed') + "careName reports that problems have " + nvar('rapidWorsen') + "rapidly worsened in the last two weeks.  ",
                'sectionTitle' : 'Cognitive and Behavioral Review of Systems',
                'MemDesctitle' : 'Memory',
                'MemDescOutput' : nvar("memLimit") + nvar('ecogMemList') + nvar('memStart'),
                'execDescTitle' : "Executive Function",
                'execDescOutput' : str(nvar('execLimit') + nvar('ecogExecList') + nvar('execJudge') + "  " + nvar('execStart')),
                'langTitle' : "Language",
                'langDescOutput1' : nvar('langLimit') + nvar('ecogLangList') + nvar('langQual') + " ",
                'langDescOutput2' : "There has been some change in " + buildlist('langElseList') if buildlist('langElseList') != "" else
                    "",
                'langDescOutput3' : "  " + nvar('langStart'),
                'visuospTitle' : "Visuospatial",
                'visuoSpoutput' : svar('visuospLimit') + "  " + nvar('ecogVisuospList') + nvar('visuospFamPeopleObj') + "  " + nvar('visuospStart'),
                'moveDescTitle ' : "Motor",
                'moveDesc1' : " He/she has " + nvar('fallYN') + "had significant falls in the past.  Fall precautions are " + nvar('fallCaut') + "in place in the home.  " + "He/she has " + nvar('weak') + "sensation of physical weakness.  There is " + nvar('fineMot') + "difficulty using his/her hands.  "+ "He/she has " + nvar('involMot1') +"noted involuntary movements" + nvar('involMot2') +". " + nvar('moveInvolMovSev') +" "+ nvar('swallow'),
                # ~ 'moveDesc2' : "" if svar('moveHyg') == "" and nvar('moveHobby') == "" and nvar('moveTurn') == "" and nvar('moveGetUp') == "" and nvar('moveFreeze') == "" else "In addition, movement problems cause " + nvar('hygMov') + "problems with hygiene, " + nvar('hobMov') + "problems with hobbies, " + nvar('turnMov') + "problems turning in bed, " + nvar('getUpMov') + "trouble getting up, and " + nvar('freezeMov') + "trouble with freezing gait. ",
                'moveDesc3' : nvar('slownessInit') + nvar('rigidityMotion') + nvar('lossPost') + nvar('restTremor'),
                'sensTitle' : "Sensory",
                'senseOutput' : svar('senseSev') + "  He/she " + nvar('numbSens') + " numbness and tingling.  He/she has" + nvar('hear') + "hearing problems.  " +  "There are also " + nvar('sight1') + "vision problems" + nvar('sight2') + "  " + nvar('smell'),
                'autoTitle' : "Autonomic",
                'autoOutput' : svar('autoFaint') + "  " + svar('autoFatigue') + "  " + svar('autoUrinIncont') + "  " + "There is " + nvar('bowel') + " incontinence of bowel.",
                'behDescTitle' : "Behavioral",
                'behDesc1' : nvar('behLimit') + "The NPI-Q (Kaufer et al. J Neuropsychiatry Clin Neurosci 2000) was filled. " + nvar('npiqList'),
                'behDesc2' : "He/she has "+ nvar('labile') + "been more emotionally labile.  " + nvar('compulse'),
                'behList' : "There is " + nvar('behInapprop') + "report of socially inappropriate behaviors, " + nvar('behOralFix') + 'report of fascination with putting things in the mouth, ' + nvar('behReligChange') + "report of changes in religious practice, " + nvar('behIntimacyChange') + "report of changes in personal intimacy, and " + nvar('behHygieneChange') + "report of changes in personal hygiene.  ",
                'behDesc3' : "He/she " + nvar('passSI') + "thoughts that he/she would be better off dead or wished he/she could go to sleep and not wake up. " + nvar('actSI2') + " he/she " + nvar('actSI1') + "thoughts of killing him/herself.  " + "He/she "+ nvar('violent') + "threatening, violent or had serious thoughts of harming others. ",
                'behDesc4' : "The patient " + nvar('readEmoEyes') + "able to correctly read people's true emotions through the eyes.  In conversations the patient " + nvar('detectFaceExp') + "sensitive to even the slightest change in the facial expression of the person he/she is conversing with.  The patient's powers of intuition " + nvar('understandingOthers') + "good when it comes to understanding others.  The patient " + nvar('detectJokeTaste') + "tell when others consider a joke in bad taste, even though they may laugh convincingly.  The patient " + nvar('detectInappropriate') + "usually tell when he/she said something inappropriate by reading it in the listener's eyes.  If someone is lying to the patient, he/she " + nvar('detectLying') + "it at once from that person's manner or expression.  ",
                'behDesc5' : nvar('behStart'),
                'sleepTitle' : "Sleep",
                'sleepDesc1' : nvar('sleepInsomnia') + "  " + nvar('sleepEDS') + "  " + nvar('epwFiller') + nvar('epworthList') + nvar('drowsy') +  "During the day, he/she is " + svar('sleepDiffArouse').lower() + ".  ",
                'sleepNapDesc' : "He/she doesn't usually nap during the day, " if contains_response(['sleepNap'], ["Not at all", 'not at all']) == True else
                    "He/she usually naps " + nvar('sleepNap') + " during the day, " if not contains_response(['sleepNap'], ["Not at all", 'not at all']) == True else
                    "",
                'sleepDesc2' : nvar('snore') + "is described, " + nvar('osa1') + "there is " + nvar('osa2') + "observed sleep apnea.  " + nvar('parasom') + nvar('rsbd') + "acting out of dreams is described (Postuma Mov Dis 2012). ",
                'faqTitle' : "Functional Independence",
                'faqOutput' : nvar('faqList') + nvar('adlList') + "He/she is " + nvar('drive') + "driving.  There have been " + nvar('mva') + "accidents. There are " + nvar('safeHome') + "safety concerns at home. ",
                'careDescTitle' : "Caregiver Stress",
                'careDescOutput' : "No information from the caregiver was available." if score('ptfillYN') == 0 else nvar('careDesc1'),
                # ~ 'careDescOutput2' : "A form about caregiver burden was offered but not filled." if svar('zarit') == '[]' else
                'careDescOutput2' : "A form about caregiver burden was offered but not filled." if surveyscore('zarit') == '[]' else
                    svar('ptId') +"'s " + svar('careRelate').lower() + ", " + svar('careName') + " also answered specific questions about caregiver burden (from Zarit, Orr, and Zarit 1985).  This suggested the following: " + nvar('zaritList') if len(nvar('zaritList')) !=0 else
                    "A form about caregiver burden was offered but not filled."},

            'medRos' : "Medical Review of Systems: A complete review of systems was done and is otherwise unremarkable except as mentioned above.",

            'probListDesc' : "Of particular interest is " + str(svar('probList')) + ". " + '\n' if svar('probList') != "" and svar('probList') != "None of the above." else '',

            'pmh' : { # past medical history
                'title' : "Past Medical History: The past medical history was directly entered into EPIC by the patient or caregiver as: *** . ",
                'staringSpells' : "There have been staring spells.  " if svar('staringSpells') == 'Yes' else
                    'There have been potential staring spells' if svar('staringSpells') == "Maybe" else
                    "",
                'illogicalThinking' : "careName reports that there has been illogical thinking.  "  if svar('illogicalThinking') == "Yes" else
                    'careName reports that there has been potential illogical thinking.  ' if svar('illogicalThinking') == "Maybe" else
                    ""},

            'psh' : "Past Surgical History: The past surgical history was directly entered into EPIC by the patient or caregiver as: *** . ",

            'allergy' : "Allergies: Allergies were directly entered into EPIC by the patient or caregiver as follows: ***.",

            'meds' : "Medications: Medications were directly entered into EPIC by the patient or caregiver, and verified as follows: *** .",

            'famHxDesc' : "Of particular interest is a family history of " + str(svar('famProbList')) + ". " if svar('famProbList') != "" else "",

            'famHx' : {
                'title' : "Family History: The family history was directly entered into EPIC by the patient or caregiver as: *** ",
                'famProbListDesc' : 'careName reports a family history of ' + str(svar('famProbList')) + ".  " if svar('famProbList') != "" else
                    ""},

            'socHxTitle' : "Social History",
            'socHx1' : nvar('dM1'),
            'socHx2' : "He/she was educated through the level of " + nvar('highestEduc') + ".  He/she received " + nvar('gradeType') + " in school, and required " + nvar('specEd') + "additional educational services or accommodations.  ",
            'socHx3' : "He/she primarily worked as " + an(svar('profType')) + " " + svar('profType') + ".  ",
            'socHx4' : " He/she now lists housemates as: " + nvar('houseMate') + ".  " + "He/she has " + nvar('nearbyFriends') + "local friends, and is " + nvar('social') + "generally social.  ",
            'socHx5' : "He/she currently exercises about " + svar('exerciseHours').lower() + " weekly.  "  + nvar('diet') + nvar('tox') + nvar('gun'),
            'outputOutro' : '',

            'obj' : "OBJECTIVE: The physical examination, labs, prior procedures and imaging results can be entered here *** ",

            'initAssess' : {
                'title' : "Preliminary Assessment",
                'output' : nvar('domCC') + nvar('fxnSummary') + nvar('fxnSum2'),
                'demRisk' : "Potential factors contributing to changing mood or cognition include: " + str(buildlist('demRiskList')) if buildlist('demRiskList') != "" and  buildlist('demRiskList') != "!!!___demRiskList___!!!" else
                    "",
                'uhoh' : "Potentially dangerous situations could stem from the following: " + str(buildlist('dangerList')) if buildlist('dangerList') != "" and buildlist('dangerList') != "!!!___dangerList___!!!" else
                    "",
                'output3' : nvar('careBurden') + nvar('NPIQCareBurden')},

            'initPlan' : {
                'initPlantitle' : "Initial Plan",
                'workup' : "Ensure B12 and TSH have been checked.  Consider checking HIV and RPR. \nEnsure we have a copy of neuroimaging to review. \nIf not done, consider detailed neuropsychological evaluation.",
                'demRiskPlan' : {
                    'vascRisk' : "The need to attend to any vascular risk factors was discussed.",
                    'depAnx' :  "Treatment of any depression or anxiety symptoms, including speaking with a qualified and trained professional as well as consideration of an antidepressant, was discussed." if svar('behDepress') != "Not present" or svar('behAnx') != "Not present" else
                        "",
                    'sleepHyg' : "Guidance on proper sleep hygiene was provided." if svar('sleepEDS') != 'No daytime sleepiness' or surveyscore('epworth') > 9 else
                        "",
                    'sleepSpecialist' : "Consultation of a sleep specialist was considered." if svar('sleepEDS') != "No daytime sleepiness." or surveyscore('epworth') > 15 else
                        "",
                    'NoctPulseOx' : "Nocturnal pulse oximetry was discussed." if svar('sleepOSA') != "No" or svar('sleepSnore') == "Yes" else
                        ""},

                'safety' : {

                    'gunStored' : "Ensure that guns are well secured." if svar('gunSecure') != "Yes" and svar('gunNum') != "none" else
                        "",
                    'violentBeh' : "Management of violent behavior was discussed, including a need to emphasize personal safety and when to involve police if necessary. \nPharmacological management of violent behavior was considered. \nNonpharmacological management of inappropriate behavior was also discussed." if svar('behViolent') != "No" else
                        "",
                    'suic' : "Active suicidal ideation was discussed.  A prevention plan was enacted, and information on suicidal ideation provided." if svar('behActSI') != "No" else
                        "",
                    'passSuic' : "" if svar('behPassiveSI') != "No" else
                        "",
                    'driving' : "Concerns regarding driving were discussed.  A driving evaluation was recommended." if (svar('adlDrive') == "Yes" and svar('adlMVA') == "Yes") or (surveyscore('ecogVisuosp') > 2 and svar('adlDrive') == "Yes") or (surveyscore('FAQ') > 6 and svar('adlDrive') == "Yes") else
                        "",
                    'homeSafety' : "Concerns regarding safety at home were discussed, as well as strategies to maximize both safety and independence." if nvar('safeHome') != "no " else
                        "",
                    # moveWalkBal doesn't exist
                    'fallRisk' : "The risks of falls were discussed.  PT/OT and home fall precautions were considered." if (svar('moveFall') != "No" or svar('moveWalkBal') != "Not at all (no problems).") and svar('moveFallCaution') != "Yes" else
                        "",
                    'safeSwallow' : "A swallow evaluation was considered." if svar('moveSwallow') != "There are no problems." else
                        "",
                    'speechTher' : "Speech therapy was considered." if surveyscore('ecogLang') > 20 else
                            ""},

                'planning' : {

                    'advanDir' : "We discussed the importance of discussing advanced directives and power of attorney before these are really needed.  The need for proactive planning, including financial, legal and care planning was discussed.",
                    'lifePlanning' : "End of life planning including screening, treatment and placement was discussed. We discussed alternative care options available in the community. We discussed the need to explore care options well before there was a crisis. \nThe anticipated need for greater family involvement and support was discussed, as well as the need to recognize family limitations.\nThe idea of palliative or supportive care consults were introduced. \nRecognition of grieving process for the family members was discussed.  We recommended they explore what kinds of financial assistance for which the patient might qualify." if surveyscore('FAQ') > 15 else
                        "",
                    'hospicePlanning' : "The idea of hospice was introduced. \nAutopsy was discussed. \nWe discussed bereavement services." if surveyscore('FAQ') > 25 else
                        ""},

                'genHealth' : "The importance of a healthy diet was emphasized, as was the need for regular physical exercise (if discussed with primary care), mental activity and social activity. \nOpportunities to participate in research were discussed.",

                'research' : "The patient was informed about available research studies through the University of Colorado in case that they are interested in participating.",

                'caregiver' : {

                    'caregiverHealth' : "The need for all caregivers involved to also attend to their own health was emphasized.  We recommend each caregiver having at least one day a week to do something they enjoy on their own.  Caregiver groups were also discussed.  Support groups were recommended. \nThe need for emergency plans was discussed in the event of caregiver illness.",
                    'socialWork' : "Additional resources such as social work were offered given signs of caregiver burnout." if surveyscore('zarit') >= 9 else
                        ""},

                'scorebreak' : "",

                'scores' : {
                    'ecogMemory':"ECog Memory Total: " + str(surveyscore('ecogMem')) +" (out of 32 possible, " + str(surveypercent('ecogMem')) + " % of maximum possible score)",
                    'ecogExec':"ECog Executive Function Total: " + str(surveyscore('ecogExec')) + " (out of 48, " + str(surveypercent('ecogExec')) + " % of maximum possible score)",
                    'ecogLang':"ECog Language Function Total: " + str(surveyscore('ecogLang')) + " (out of 27, " + str(surveypercent('ecogLang')) + " % of maximum possible score)",
                    'ecogVisuo':"ECog Visuospatial Function Total: " + str(surveyscore('ecogVisuosp')) + " (out of 21, " + str(surveypercent('ecogVisuosp')) + " % of maximum possible score)",
                    'npqiSeverity':"NPI-Q Severity Total: " + str(surveyscore('npiq')) + " (out of 36, " + str(surveypercent('npiq')) + " % of maximum possible score)",
                    'npqiCaregiver':"NPI-Q Caregiver Burden: "+ str(surveyscore('npiqCaregiver')) + " (out of 60, " + str(surveypercent('npiqCaregiver')) + " % of maximum possible score)",
                    #'ecogMemory':"Cognitive Fluctuation Score: " + str(compositescore('Flux')) +" (out of 4 possible, " + str(compositepercent('Flux')) + " % of maximum possible score)",
                    'epworth':"Epworth Score: " + str(surveyscore('epworth')) + " (out of 24, " + str(surveypercent('epworth')) + " % of maximum possible score)",
                    'FAQ':"Functional Activities Questionnaire Total: " + str(surveyscore('FAQ')) + " (out of 30 possible, " + str(surveypercent('FAQ')) + " % of maximum possible score)",
                    'Zarit':"Zarit Caregiver Burden: " + str(surveyscore('zarit')) +" (out of 48 possible, " + str(surveypercent('zarit')) + " % of maximum possible score)",
                }
            }
        }

          #TESTING SECTION TO TEST AND TRY TO BREAK NESTING LOGIC, FUNCTIONS ETC.
          # 'testingsection' : {

          #     'outputvariable1' : "HEY, I'M VARIABLE 1!!",
          #     'outputvariable2' : "HEY, I'M VARIABLE 2!!",
          #     'outputvariable3' : "HEY, I'M VARIABLE 3!!",
          #     'nestedvariable1' : {'subvariable1' : "Hey, I'm subvariable 1.",
          #             'subvariable2' : "Hey, I'm subvariable 2.",
          #             'subvariable3' : {
          #                 'subsubvariable1' : "Hey, I'm sub-subvariable 1.",
          #                 'subsubvariable2' : "Hey, I'm sub-subvariable 2.",
          #                 'subsubvariable3' : {'subsubsubvariable1' : "Hey, I'm sub-sub-subvariable1",
          #                     'subsubsubvariable2' : {
          #                         'title' : "Hey, here is the title of sub-sub-subvariable1",
          #                         'variables' : {
          #                             'sub-sub-subvariable1variable1' : "hey, I am really, really nested!",
          #                             'sub-sub-subvariablevariable2' : "hey, I am also really really nested!!"},
          #                             'sub-sub-subvariablevariable3' : {
          #                                 'toolong' : 'Whoopie!',
          #                                 'toolong2' : "Hootie Hoo!",
          #                                 'toolong3' : {"blah blah blah" : 'super nest!!!',
          #                                     'supernest2' : 'wheeeeeeeeeeeeeeeee'}},
          #                         'output' : nvar('sub-sub-subvariable1variable1') + nvar('sub-sub-subvariablevariable2') + nvar('sub-sub-subvariablevariable3')}}}}},






    #takes a string containing the name of a key you want to find
    #function that recursively searches obj until it finds the key in question, and returns the value of the key
    #watch out for keys that appear more than once (i.e. 'output')!
    #STP: This is a icky way of getting data. I think it would be cleaner to
    # require the user to actually know how to find the data.
    def findvar(self, obj, key):
        if key in obj:
            return obj[key]
        # STP: Might need to catch this in case obj is not a dict
        for v in obj.values():
            if isinstance(v, dict):
                item = self.findvar(v, key)
                # STP: Is 'not none' correct, or does item just need to be truthy?
                if item is not None:
                    return item

    #takes a top-level dictionary (i.e. self.note) and a key (i.e. 'execStart') as input, and returns the full path of the key
    #i.e. "self.note['hpi']['memDesc']['execStart']"
    #STP: Unused
    # ~ def findpath(self, obj, key):
        # ~ dictionary = eval(obj)
        # ~ for k,v in dictionary.items():
            # ~ if k == key:
                # ~ return obj + "['"+k+"']"
            # ~ elif isinstance(v,dict):
                # ~ out = self.findpath(obj + "['"+k+"']", key)
                # ~ if out is not None:
                    # ~ return out

    #function that list of keys in a dictionary, returns the value of each key.  If the key's value is a dictionary (i.e. a nested dictionary), this loops
    #through that dictionary and returns all of those values as well.  Used to generate the text of your note.
    # old and unused
    def write(self, key):
        value = self.findvar(self.note, str(key))
        if isinstance(value, dict):
            mystring = ""
            for k in value:
                mystring += self.write(k)
            return mystring
        else:
            return value


    #short for "note variable", returns the value in self.variables with the key you provide as input.
    #Make sure to name each key differently, as this won't work if two keys have the same name.
    #If the variable can't be found, returns the below pattern with the variable name on the inside.
    def nvar(self, key):
        #STP: Is 'not none' the right check here or does it have to be truthy?
        if self.findvar(self.variables, str(key)) is not None:
            return self.findvar(self.variables, str(key))
        else:
            print("nvar can't find " + key + "!!!", file=sys.stderr)
            print(format_stacktrace(), file=sys.stderr)
            return " nvar(###___" + key + "___###) "



    #short for "survey variable", returns the value in self.survey that you provide as input.  If the question type is multiple response,
    #returns a list separated by commas.  If the variable can't be found (for example if it's typed wrong in your survey class or you took the question out of the survey),
    #this function returns the name of the variable, with a hashtag on either side (i.e. #variable1#.)
    def svar(self, variable):
        if variable in svar_blacklist:
            print(f"Not looking for variable {variable} because it is on the svar blacklist")
            return f' svar(#{variable}#) '

        #checks to make sure the variable isn't within a standardized survey.  If it is, tries to return the value here.
        surveyvariableslist = []
        for survey in self.standardizedsurveys.values():
            for ii in survey[0]:
                surveyvariableslist.append(ii)

        if variable not in surveyvariableslist:
            try:
                try:
                    #if this question can have multiple answers selected
                    if self.survey.sections[self.survey.section(str(variable))][str(variable)][1] == 4:
                        #and if it's not empty
                        if self.survey.__getattribute__(str(variable)) != []:
                            #if there's only one answer, just return that
                            if len(self.survey.__getattribute__(str(variable))) == 1:
                                return self.survey.__getattribute__(str(variable))[0]
                            #otherwise, make it into a readable sentence separated by commas and return the sentence
                            else:
                                allbutlast = ', '.join(self.survey.__getattribute__(str(variable))[:-1])
                                last = self.survey.__getattribute__(str(variable))[-1]
                                itemlist = ', and '.join([allbutlast, last])
                                return itemlist
                        else:
                            #if it's not a question that can have multiple responses, there must only be one response, so just return that.
                            return self.survey.__getattribute__(str(variable))
                except KeyError as e:
                    print("svar can't find variable " + str(variable) + "!", file=sys.stderr)
                    print(format_stacktrace(), file=sys.stderr)
                    return f' svar(#{variable}#) '
                return self.survey.__getattribute__(str(variable))
            except AttributeError as e:
                print("svar can't find variable " + str(variable) + "!", file=sys.stderr)
                print(format_stacktrace(), file=sys.stderr)
                return f' svar(#{variable}#) '
        else:
            try:
                return self.survey.__getattribute__(str(variable))
            except AttributeError as e:
                print("svar can't find variable " + str(variable) + "!", file=sys.stderr)
                print(format_stacktrace(), file=sys.stderr)
                return f' svar(#{variable}#) '

    #takes a function with parameters (for example, svar('ptAge')) and the expected null output (i.e. !!!___demRiskList___!!! if this variable is not present, and
    #returns an empty string if the null output is produced by the function.  Handy for questions or surveys that are only displayed if other questions are answered in a certain way.
    def conditional(self, function, expectednulloutput):
        outstring = function
        if outstring == expectednulloutput:
            return ""
        else:
            return outstring

    #takes a list of variables and a list of responses (['adlDrive', 'adlMVA'], ['not distressing at all', 'very distressing']) as input,
    #and returns true if any of the variables are assigned to any values of the list.  Note that this function does NOT return "False" if it is not present,
    # so if you want to test a negative, use: "if contains_response(variables, responses) is not True:"
    def contains_response(self, variables, responses):
        for x,y in zip(variables, cycle(responses)) if len(variables) > len(responses) else zip(cycle(variables), responses):
            try:
                if y in self.svar(str(x)):
                    return True; break
            except TypeError:
                #TODO
                print("contains_response can't find one or more variables!", file=sys.stderr)


    #takes the name of a survey and types of survey responses (i.e. 'FAQmod', 'FAQsev') as input, returns true if there were any survey responses in this category,
    #else returns False
    def anysurveyresponses(self, surveyname, *args):
        # STP: What is a better name for i?
        for i in self.survey.sections[str(self.survey.section(surveyname))][surveyname][0]:
            # STP: What is a better name for ii? 'arg'?
            for ii in args:
                if self.svar(i) == self.survey.sections[str(self.survey.section(surveyname))][surveyname][2][ii][0]:
                    return True
        return False

    #takes a dictionary of keys and values as an input, formats it in a list sentence structure (i.e. values xyz are turned into "x, y, and z")
    # STP: I don't think this should append the ". ", which makes it only useful
    # at the end of the sentence
    def buildlist(self, key):
        itemlist = []
        listout = []
        try:
            # STP: What is a better name for i?
            for i in self.findvar(self.variables, key):
                #this function only works on flat dictionaries with no nesting in them.
                if isinstance(self.findvar(self.variables, str(i)), dict):
                    return "Too much dictionary nesting in variable " + key + " for buildlist!!"
                else:
                    #if the value of the variable in the note is not "", add it to the list
                    if self.findvar(self.variables, i) != '':
                        itemlist.append(self.findvar(self.variables, i))
            if itemlist != []:
                if len(itemlist) == 1:
                    itemlist = str((str(itemlist[0]) + ". "))
                    return itemlist
                if len(itemlist) == 2:
                    itemlist = str(str(itemlist[0]) + " and " + str(itemlist[1]) + ". ")
                    return itemlist
                else:
                    allbutlast = ', '.join(itemlist[:-1])
                    last = itemlist[-1]
                    itemlist = ', and '.join([allbutlast, last])
                    itemlist = itemlist + ". "
                    itemlist = ''.join(itemlist)#STP: I think this is unnecessary...
                    return itemlist
                    #STP: This does the same thing, but is harder to read perhaps?
                    # return ', '.join(itemlist[:-1]) + ", and " + itemlist[-1] + ". "
            else:
                return ""
        except TypeError:
            print("buildlist can't find variable " + key, file=sys.stderr)
            return f" buildlist(!!!___{key}___!!!) "

    #builds a paragraph description of a standardized survey by mapping the value of the standardized survey response to the "standardizedsurveys" dictionary in this class.
    def buildsurveydesc(self, surveyname):
        listout = []
        surveyname = str(surveyname)
        # ~ print(f"Working with survey description {surveyname}")
        #STP: Should use possibleresponse, sentencebeginning in X.items()
        for possibleresponse in self.standardizedsurveys[surveyname][1]:
            # ~ print(f"Possible response {possibleresponse}")
            itemlist = []
            sentencebeginning = self.standardizedsurveys[surveyname][1][str(possibleresponse)]
            for variable in self.standardizedsurveys[surveyname][0]:
                # ~ print(f"Possible variable {variable}")
                try:
                    if self.survey.__getattribute__(str(variable)) == self.survey.sections[self.survey.section(surveyname)][surveyname][2][str(possibleresponse)][0]:
                        itemlist.append(self.standardizedsurveys[surveyname][0][str(variable)])
                        # ~ print(f"Adding variable {variable}")
                except AttributeError:
                    print("buildsurveydesc can't find variable " + variable + "!", file=sys.stderr)
                except KeyError:
                    print("buildsurveydesc can't find survey " + surveyname + "!", file=sys.stderr)
            if itemlist != []:
                if len(itemlist) == 1:
                    listout.append(str(sentencebeginning) + str(itemlist[0]) + ".  ")
                else:
                    allbutlast = ', '.join(itemlist[:-1])
                    last = itemlist[-1]
                    itemlist = ', and '.join([allbutlast, last])
                    listout.append(str(sentencebeginning) + str(itemlist) + ".  ")
        listout = ''.join(listout)
        # ~ print(f"Got list out {listout}")
        return listout

    #builds a dictionary of a standardized survey by mapping the value of the standardized survey response to the "standardizedsurveys" dictionary in this class.
    def buildsurvey_dict(self, surveyname):
        dictionary = {"Question":"Response"}
        surveyname = str(surveyname)
        # ~ print(f"Working with survey dict {surveyname}")
        for variable, value in self.standardizedsurveys[surveyname][0].items():
            # ~ print(f"Possible variable {variable}")
            try:
                dictionary[value.capitalize()] = self.survey.__getattribute__(str(variable))
                # ~ print(f"Adding variable {variable}")
            except AttributeError:
                print("buildsurveydesc can't find variable " + variable + "!", file=sys.stderr)
            except KeyError:
                print("buildsurveydesc can't find survey " + surveyname + "!", file=sys.stderr)
        # ~ print(f"Got this dict {dictionary}")
        return dictionary

    #takes a list of names of surveys as input, and returns the one with the highest percentage.
    def domain(self, *surveys):
        arglist = [(self.surveypercent(i),i) for i in surveys]
        try:
            domCompList = sorted(arglist)
            return domCompList[-1][1]
        except Exception as e:
            print(f"Sorting didn't work!  Check your variables. {e}", file=sys.stderr)

    #returns the numerical value associated with a response.
    def score(self, variable):
        return self.survey.responsenumber(variable)

    #simplifies a score on a 0-4 scale to a 0-1 scale, either returning 0, 0.5, or 1
    def score0to1(self, variable):
        score = self.score(variable)
        if score > 2:
            score = 1
        elif score > 0 and score <= 2:
            score = 0.5
        else:
            score = 0
        return score

    #returns a raw total score of a survey.  This will depend on how many possible responses there are, etc., so it's contextual based upon the survey being scored.
    def surveyscore(self, surveyname):
        surveysection = self.survey.section(surveyname)
        score = 0
        try:
            for x in self.survey.sections[surveysection][surveyname][0]:
                score += self.survey.responsenumber(x, surveyname)
            return score
        except KeyError:
            #STP: Might want to use logging here rather than just printing.
            print("surveyscore can't find the " + surveyname + " survey!", file=sys.stderr)
            print(format_stacktrace(), file=sys.stderr)

            return 999

    #returns a survey score as a percentage, based upon how many questions there are, and what the possible total number of points is
    def surveypercent(self, surveyname):
        try:
            score = self.surveyscore(surveyname)
            numberofquestions = len(self.survey.sections[self.survey.section(surveyname)][surveyname][0])
            possiblepoints = numberofquestions * self.survey.highestresponse(surveyname)
            return round((score / possiblepoints * 100), 2)
        except KeyError:
            print("surveypercent can't find " + surveyname + "!", file=sys.stderr)

    #will sum the score of many different individual scores.  Takes a list of strings of variable names as input.
    def compositescore(self, variables):
        score = 0
        for x in variables:
            score = score + self.score(str(x))
        return score

    #returns the percentage for the sum of different individual scores.  NOTE: assumes that each question has the same possible points
    #or that if they are different that this doesn't matter.
    def compositepercent(self, variables):
        score = self.compositescore(variables)
        numberofvariables = len(variables)
        try:
            return round((score / numberofvariables * 100), 2)
        except ZeroDivisionError:
            print("Compositepercent couldn't find any variables!", file=sys.stderr)

    #uses the replace function to replace his/her, he/she etc. with the patient's preferred gender pronoun.
    def genderpronouns(self, body):
        if self.svar('ptSex') == "Male":
            body = body.replace('his/her', 'his')
            body = body.replace('His/her', 'His')
            body = body.replace('him/her', 'him')
            body = body.replace('Him/her', "Him")
            body = body.replace('he/she', 'he')
            body = body.replace('He/she', "He")
        if self.svar('ptSex') == "Female":
            body = body.replace('his/her', 'her')
            body = body.replace('His/her', 'Her')
            body = body.replace('him/her', 'her')
            body = body.replace('Him/her', "Her")
            body = body.replace('he/she', 'she')
            body = body.replace('He/she', "She")
        if self.svar('ptSex') == "Other":
            body = body.replace('his/her', 'their')
            body = body.replace('His/her', 'Their')
            body = body.replace('him/her', 'them')
            body = body.replace('Him/her', "Them")
            body = body.replace('he/she', 'they')
            body = body.replace('He/she', "They")
        return body

    #takes a list of variables as input, and prints the question for each, a colon, and the response, followed by a new line.
    def questionsandresponses(self, variables):
        outstring = ""
        for x in variables:
            section = self.survey.section(x)
            newline = str(self.survey.sections[str(section)][str(x)][0] + ":  " + self.svar(x) + '\n')
            outstring = outstring + newline
        return outstring

    #returns "a" or "an", depending on what the next word is (sometimes needed when the next word is a dynamic variable)
    def an(self, nextword):
        if nextword.startswith("a") or nextword.startswith("e") or nextword.startswith("i") or nextword.startswith("o") or nextword.startswith("u") or nextword.startswith("8") == True:
            return "an"
        else:
            return "a"

    #returns the value of the function (i.e. svar('myvariable') with a space on either end if a condition is true (i.e. svar('variable') == 'Yes'), otherwise returns a " ".
    #Useful for when you want to either put some value in the middle of a sentence (which will need two spaces, one on on either side),
    #or put nothing (which will only be one space).
    def space(self, function, condition):
        if condition == True:
            return " " + function + " "
        else:
            return " "

    def table(self, survey):
        pass

    def export_to_text(self, filename):
        whole_note = self.add_dict_to_text(self.note)

        #Replaces he/she, his/her etc. with the correct gender pronoun
        whole_note = self.genderpronouns(whole_note)

        #replaces the "Carename" variable with the care provider's name
        whole_note = whole_note.replace('careName', self.svar('careName'))

        print(f'Saving text note to {filename}')
        with open(filename, 'w') as f:
            f.write(whole_note)

    # This removes the findvar completely
    def add_dict_to_text(self, dictionary, text=""):
        for key, value in dictionary.items():
            if 'table' in key.lower():
                # ~ print("Found a table!")
                text = self.add_table_to_text(value, text)
            elif isinstance(value, dict):
                text = self.add_dict_to_text(value, text)
            elif 'title' in key.lower(): # Check for titles in case insensitive way
                # ~ print("Found a title!")
                value = value.strip()
                if value:
                    text += "\n" + self.clean_heading(value)
                    text += "\n=========================================\n\n"
            else:
                value = self.genderpronouns(value).replace('careName', self.svar('careName'))
                value = self.clean_body(value)
                if '\n' in value:
                    for v in value.split('\n'):
                        v = v.strip()
                        if v: # remove empty paragraphs
                            # ~ d.add_paragraph(v)
                            text += v + "\n\n"
                else:
                    if value:
                        text += value + "\n\n"
        return text

    def add_table_to_text(self, table, text=''):
        for key, value in table.items():
            if isinstance(value, dict):
                print(f"Dictionaries are not allowed in tables. key: {key}, table:{table}")
            elif 'title' in key.lower(): # Check for titles in case insensitive way
                print(f"Titles are not allowed in tables. key: {key}, table:{table}")
            elif 'table' in key.lower():
                print(f"Tables are not allowed in tables. key: {key}, table:{table}")
            else:
                text += key + ": "
                if '\n' in value:
                    for v in value.split('\n'):
                        v = v.strip()
                        if v: # remove empty paragraphs
                            # ~ d.add_paragraph(v)
                            text += v + "\n\n"
                else:
                    text += value + "\n\n"
        return text


    def export_to_docx(self, filename):
        d = Document()
        self.add_dict_to_document(self.note, d)

        f = d.styles['Heading 1'].font
        f.color.rgb = RGBColor(0xCF,0xB8,0x7C)
        f.highlight_color = WD_COLOR_INDEX.BLACK

        print(f'Saving Word doc to {filename}')
        d.save(filename)


    def add_dict_to_document(self, dictionary, d):
        for key, value in dictionary.items():
            if 'table' in key.lower():
                print("Adding table to document!")
                self.add_table_to_document(value, d)
            elif isinstance(value, dict):
                self.add_dict_to_document(value, d)
            elif 'title' in key.lower(): # Check for titles in case insensitive way
                if value:
                    d.add_heading(self.clean_heading(value))
                # ew
                # ~ t = d.add_table(1, 1)
                # ~ t.cell(0,0).text = self.clean_heading(value)
            else:
                value = self.genderpronouns(value).replace('careName', self.svar('careName'))
                value = self.clean_body(value)
                if '\n' in value:
                    for v in value.split('\n'):
                        v = v.strip()
                        if v: # remove empty paragraphs
                            d.add_paragraph(v)
                else:
                    if value:
                        d.add_paragraph(value)

    def add_table_to_document(self, table, d):
        t = d.add_table(0,2)
        for key, value in table.items():
            if isinstance(value, dict):
                print(f"Dictionaries are not allowed in tables. key: {key}, table:{table}")
            elif 'title' in key.lower(): # Check for titles in case insensitive way
                print(f"Titles are not allowed in tables. key: {key}, table:{table}")
            elif 'table' in key.lower():
                print(f"Tables are not allowed in tables. key: {key}, table:{table}")
            else:
                r = t.add_row()
                r.cells[0].text = key
                value = self.genderpronouns(value).replace('careName', self.svar('careName'))
                value = self.clean_body(value)
                r.cells[1].text = value
                # ~ if '\n' in value:
                    # ~ for v in value.split('\n'):
                        # ~ v = v.strip()
                        # ~ if v: # remove empty paragraphs
                            # ~ d.add_paragraph(v)
                # ~ else:
                    # ~ d.add_paragraph(value)



    def clean_heading(self, heading):
        if heading[-1] == ':': # Remove the trailing colon
            heading = heading[:-1]
        heading = heading.strip()
        heading = heading.replace('\n', ' ')
        return heading

    def clean_body(self, body):
        while '\n\n' in body:
            body = body.replace('\n\n', '\n')

        return body

    #will be a function that returns true if the variable has been filled out
    def iscomplete(self, question):
        pass

    #will output all variables in a particular section (instead of having to to 'output' : nvar('variable1') + nvar('variable2') + nvar('variable3')....etc)
    def allvariables(self, section):
        pass




    #OUTDATED, HERE FOR STORAGE.
    #This function searches through the text for a particular pattern (right now it is set to the pattern generated for when an nvar() cannot find a variable (i.e. "###___variable___###"))
    # and re-initializes any variables that match this pattern.  This issue was fixed by restructuring the note, but this function has been left here in case it is useful in the future
    # for searching the text and performing a certain operation on pieces of it.
    def replacenvars(self, body):
        for string in body.split():
            try:
                m = re.search('###___(.+?)___###', string)
                if m is not None:
                    index = body.index(m.group(0))
                    lenstring = len(m.group(0))
                    #takes out the extra space in front of the ###___
                    if body[index+lenstring] == " ":
                        body = body[:index+lenstring] + body[index+lenstring+1:]
                    #takes out the extra space after the ###___
                    if body[index-1] == " ":
                        body = body[:index-1] + body[index:]
                    #replaces the ###___variable___### with the value of nvar(variable)
                    body = body.replace(string, self.nvar(m.group(1)))
            except TypeError:
                pass
            except ValueError:
                print("You seem to be referencing a note variable that doesn't actually ever get defined: " + m.group(0), file=sys.stderr)
        return body

    def replacebuildlists(self, body):
        for string in body.split():
            try:
                m = re.search('!!!___(.+?)___!!!', string)
                if m is not None:
                    body = body.replace(string, self.buildlist(m.group(1)))
            except TypeError:
                pass
        return body




class Ptnote99483(Ptnote):


    def __init__(self, survey):

        self.survey = survey

        #creates an empty note, so that the variables can be initialized without bugging out
        self.note = {}
        self.whole_note = ""

        #this section just renames functions so that you don't have to type self.x() each time, you can just type x()
        findvar = self.findvar
        # ~ findpath = self.findpath
        write = self.write
        nvar = self.nvar
        svar = self.svar
        contains_response = self.contains_response
        anysurveyresponses = self.anysurveyresponses
        buildlist = self.buildlist
        buildsurveydesc = self.buildsurveydesc
        domain = self.domain
        score = self.score
        score0to1 = self.score0to1
        surveyscore = self.surveyscore
        surveypercent = self.surveypercent
        compositescore = self.compositescore
        compositepercent = self.compositepercent
        genderpronouns = self.genderpronouns
        an = self.an
        table = self.table
        # ~ export = self.export
        iscomplete = self.iscomplete
        allvariables = self.allvariables
        replacenvars = self.replacenvars
        replacebuildlists = self.replacebuildlists
        conditional = self.conditional
        questionsandresponses = self.questionsandresponses
        space = self.space

        #The DSRS is a slightly different survey, so we'll need to use the function compositescore instead of surveyscore.  This next line makes a list of all of the variables
        #in the DSRS, so you can just call compositescore(DSRS) to get the total score.
        DSRS = ['DSRSmem', 'DSRSlang', 'DSRSrecog', 'DSRSorientTime', 'DSRSorientPlace', 'DSRSdecision', 'DSRSsoc', 'DSRShomeActiv', 'DSRSpersCare', 'DSRSeating', 'DSRSurin', 'DSRSabilityPlace']

        #this section adds any standardized surveys unique to the 99483 class to the 'standardizedsurveys' dictionary that was inherited from the superclass ptnote object.
        self.standardizedsurveys['hoardingsurvey'] = ({'hoardingClutter' : 'using rooms in the home due to clutter',
            'hoardingDiscard' : 'discarding ordinary possessions',
            'hoardingCollecting' : 'collecting or buying new possessions',
            'hoardingEmoDistress' : 'emotional distress because of clutter',
            'hoardingSocImpair' : 'impairment in their daily life because of hoarding behavior'},
            {'hoardingNone0' : 'He/She reports no problems (0 on a scale of 8) with ',
            'hoardingNone1' : 'He/She reports very little problems (1 on a scale of 8) with ',
            'hoardingMild2' : 'He/She reports mild problems (2 on a scale of 8) with ',
            'hoardingMild3' : 'He/She reports mild problems (3 on a scale of 8) with ',
            'hoardingMod4' : 'He/She reports moderate problems (4 on a scale of 8) with ',
            'hoardingMod5' : 'He/She reports moderate problems (5 on a scale of 8) with ',
            'hoardingSev6' : 'He/She reports severe problems (6 on a scale of 8) with ',
            'hoardingSev7' : 'He/She reports severe problems (7 on a scale of 8) with ',
            'hoardingExt8' : 'He/She reports extreme problems (8 on a scale of 8) with '})

        self.variables = {

          'initAssessVariables' : {

              'probSum' : "no significant past medical history " if "none of the above" in svar('probList') else
                  "a past medical history of ",
              'probList' : str(svar('probList')),

              'dangerList' : {

                  'rapidWorsenDanger' : "rapid and recent decline" if svar('rapidWorsen') != "No" else
                      "",
                  'gunStore' : "firearms are not securely stored" if svar('gunSecure') != "Yes" and svar('gunHome') != 'No' else
                      "",
                  'concernViolent' : "some concern for violent behavior towards others" if svar('behViolent') != "No" else
                      "",
                  'concernActSI' : "concern for active suicidal ideation" if svar('behActSI') != "No" else
                      "",
                  'concernPassSI' : "concern for passive suicidal ideation" if svar('behPassiveSI') != "No" else
                      "",
                  'concernDriving' : "some concern regarding driving" if (svar('adlDrive') == "Yes" and svar('adlMVA') == "Yes") or (surveyscore('FAQ') > 9 and svar('adlDrive') == "Yes") else
                      "",
                  'concernSafeHome' : "concerns regarding safety at home" if svar('adlHomeSafe') != 'No' else
                      ""},

              'demRiskList' : {

                  'depRisk' : "depression" if svar('behDepress') != "Not present" else
                      "",
                  'anxRisk' : "anxiety" if svar('behAnx') != "Not present" else
                      ""}},

          'hpivariables' : {

              'careSum' : svar('careName') + ", " + svar('ptId') + "'s" + " " + svar('careRelate').lower() + ", " + "helped to provide the history.  " + svar('careName') + " sees " + svar('ptId') + " about " + svar('careFreq').lower() + " or approximately " + svar('careHr').lower() + ", is " + svar('careAge') + " years old, and has known the patient for " + svar('careYr') + " years.  ",
              'probSpeed' : 'When asked about how rapidly symptoms have changed, careName reports that they have ' + svar('probSpeed').lower() + ".  " if not contains_response(['probSpeed'], ["There has been no worsening"]) == True else
                  'When asked about how rapidly symptoms have changed, careName reports that they have not changed.  ',
              'rapidWorsen' : ' not ' if svar('rapidWorsen') == "No" else
                  " ",
              'sigFlux' : "When asked about fluctuations, "+ svar('careName') + " reports that there is " + svar('sigFluxYN') + ".  "},

              'faQvariables' : {

                  'faqList' : buildsurveydesc('FAQ'),
                  'adlList' : buildsurveydesc('ADL'),
                  'DSRS' :{'DSRStitle' : "Dementia Severity Rating Scale",
                      'DSRSmem' : 'MEMORY: ' + svar('DSRSmem'),
                      'DSRSlang' : 'LANGUAGE: ' + svar('DSRSlang'),
                      'DSRSrecog' : 'RECOGNITION OF FAMILY MEMBERS: ' + svar('DSRSrecog'),
                      'DSRSorientTime' : 'ORIENTATION TO TIME: ' + svar('DSRSorientTime'),
                      'DSRSorientPlace' : 'ORIENTATION TO PLACE: ' + svar('DSRSorientPlace'),
                      'DSRSdecision' : "ABILITY TO MAKE DECISIONS: " + svar('DSRSdecision'),
                      'DSRSsoc' : "SOCIAL AND COMMUNITY ACTIVITY: " + svar('DSRSsoc'),
                      'DSRShomeActiv' : "HOME ACTIVITIES AND RESPONSIBILITIES: " + svar('DSRShomeActiv'),
                      'DSRSpersCare' : 'PERSONAL CARE - CLEANLINESS: ' +svar('DSRSpersCare'),
                      'DSRSeating' : "EATING: " + svar("DSRSeating"),
                      'DSRSurin' : 'CONTROL OF URINATION AND BOWELS: ' + svar('DSRSurin'),
                      'DSRSabilityPlace' : 'ABILITY TO GET FROM PLACE TO PLACE: ' + svar('DSRSabilityPlace'),
                      'DSRSscore' : 'TOTAL SCORE: ???',
                      'DSRSseverity' : 'SEVERITY: ???',
                  }
              },

              'SubjSafetyVariables' : {

                  'drive' : "currently " if svar('adlDrive') == "Yes" else
                      "not ",
                  'mva' : 'possible ' if svar('adlMVA') == "Maybe" else
                      "" if svar('adlMVA') == "Yes" else
                      "no ",
                  'safeHome' : "possible " if svar('adlHomeSafe') == "Maybe" else
                      "" if svar('adlHomeSafe') == "Yes" else
                      "no ",
                  'wandering' : "There has been wandering or getting lost.  " if svar('wandering') != "No" else
                      "There has been no wandering or getting lost.  ",
                  'medications' : "Medications have been taken consistently.  " if svar('takingMeds') == "Yes" else
                      'Medications have not been taken consistently.  ',
                  'liveAlone' : "The patient is alone for a substantial portion of the day.  " if svar('liveAlone') != "No" else
                      'The patient is not alone for a substantial portion of the day.  ',
                  'fallRiskAssess' : "There is concern regarding imbalance and falls.  " if svar('unsteadyFall') != "No" else
                      "",
                  'finMismanage' : "There is concern for potential financial mismanagement.  " if svar('finMismanage') != "No" else
                      "",
                  'choking' : "There is concern regarding his/her ability to swallow safely. " if svar('choking') != "No" else
                      "",
                  'hoarding' : "There is some concern regarding clutter and/or hoarding behavior." if svar('hoarding') != "No" else
                      ""},

              'behDescVariables' : {

                  'npiqList' : buildsurveydesc('npiq'),
                  'behHygieneChange' : 'some ' if svar('behHygieneChange') == "Maybe" else 'no ' if svar('behHygieneChange') == "No" else
                  "",
                  'labile' : "" if svar('behEmoLabile') != "No" else
                      "not ",
                  'passSI' : "does not deny " if svar('behPassiveSI') != "No" else
                      "denies ",
                  'actSI1' : "does not deny " if svar('behActSI') != "No" else
                      "denies ",
                  'actSI2' : "In addition," if svar('behActSI') != "No" and svar('behPassiveSI') != "No" else
                      "However," if svar('behActSI') == "No" and svar('behPassiveSI') != "No" else
                      "Furthermore,",
                  'violent' : "is reportedly " if svar('behViolent') != "No" else
                      "also " if svar('behActSI') != "No" and svar('behViolent') != "No" else
                      "is not reportedly ",
                  'hoardingSurv' : buildsurveydesc('hoardingsurvey') if svar('hoarding') != 'No' else
                      ""},

              'careDescVariables' : {

                  'careDesc1' : svar('ptId') +"'s " + svar('careRelate').lower() + ", " + svar('careName') + ", answered several questions regarding caregiver burden and stress, starting with the NPI-Q, which suggests the following.  " + buildsurveydesc('npiqCaregiver') if buildsurveydesc('npiqCaregiver') != '' else
                      "",
                  'careDesc2' : "A form about caregiver burden was offered but not filled." if buildsurveydesc('zarit') == "" and anysurveyresponses('zarit', 'zaritNone') == False else
                      str(svar('ptId')) +"'s " + svar('careRelate').lower() + ", " + svar('careName') + " also answered specific questions about caregiver burden (from Zarit, Orr, and Zarit 1985).  This suggested the following.  " + buildsurveydesc('zarit') if anysurveyresponses('zarit', 'zaritMild', 'zaritMod', 'zaritSev', 'zaritMin') == True else
                      "" if anysurveyresponses('zarit', 'zaritNone') == True else
                      "A form about caregiver burden was offered but not filled.",
                  'needsAddressed' : 'When asked about other concerns they wanted addressed, careName reported: ' + svar('needsAddressed') if svar('needsAddressed') != "#needsAddressed#" else
                      ""
            }
        }

        self.note = {

            'initAssess' : {

                'title' : "Preliminary Assessment",
                'ptSum': str(svar('ptId')) + " is a ***-year-old " + svar('ptSex').lower() + " with " + str(nvar('probSum')) + conditional(svar('probList'), "none of the above") + " who is visiting the Rocky Mountain Memory Disorders Clinic at the University of Colorado, Denver for a comprehensive care plan.  ",
                'uhoh' : "Potentially dangerous situations could stem from the following: " + str(buildlist('dangerList')) if buildlist('dangerList') != "" and buildlist('dangerList') != "!!!___dangerList___!!!" else
                    "",

                'demRisk' : "Potential factors contributing to changing mood or cognition include: " + buildlist('demRiskList') if buildlist('demRiskList') != "" and buildlist('demRiskList') != "!!!___demRiskList___!!!" else
                    "",
                'DecisionMakingCap' : "Decision making capacity: ***",
                'FAQscoreDesc' : "Functional Activities Questionnaire Total: " + str(surveyscore('FAQ')) + " (out of 30 possible, " + str(surveypercent('FAQ')) + " % of maximum possible score)",
                'SeverityTitle' : "Level of Severity",
                'FAQdesc' : "The FAQ is suggestive of Major Neurocognitive Disorder.  " if (anysurveyresponses('FAQ', 'FAQsev') == True) or (surveyscore('FAQ') >= 9) else
                    'The FAQ is suggestive of mild cognitive impairment.  ' if surveyscore('FAQ') < 9 and surveyscore('FAQ') > 0 else
                    "The FAQ is normal.  ",
                    #Add in another here if the FAQ was not completely filled
                'DSRSdesc' : 'The DSRS suggests mild cognitive impairment.  ' if compositescore(DSRS) > 0 and compositescore(DSRS) <= 18 else
                    'The DSRS suggests moderate cognitive impairment.  ' if compositescore(DSRS) >= 19 and compositescore(DSRS) <= 36 else
                    'The DSRS suggests severe cognitive impairment.  ' if compositescore(DSRS) != 0 else
                    "The DSRS score was a 0.  ",
                'npiqDesc' : "Neuropsychiatric Symptoms: " + buildsurveydesc('npiq'),
                'npiqPlan' : "Therefore, we will ***.  " if surveyscore('npiq') > 0 else
                    "",
                'commRec' : "Community resources: ***",
                'careBurden': 'There are no obvious signs of significant caregiver burden. ' if surveyscore('zarit') == 0 else
                    'The response on the Zarit falls with between the 25th to 50th percentile. ' if surveyscore('zarit') >= 3 and surveyscore('zarit') < 9 else
                    'The response on the Zarit falls with between the 50th to 75th percentile. ' if surveyscore('zarit') >= 9 and surveyscore('zarit') < 17 else
                    'The response on the Zarit is within the top quartile for caregiver stress. ' if surveyscore('zarit') > 17 and surveyscore('zarit') != 999 else
                    "",
                'NPIQCareBurden' : "Neuropsychiatric and behavioral symptoms are significantly contributing to caregiver stress.  " if anysurveyresponses('npiqCaregiver', 'npiCareMod', 'npiCareSev', 'npiCareExt') == True else
                    "Neuropsychiatric or behavioral symptoms do not cause significant caregiver burden.  ",
                'paragraphSpace' : ''},

            'initPlan' : {

                'initPlanTitle' : "Care Plan",
                'workup' : "*** Ensure B12 and TSH have been checked. \n\n*** Consider checking HIV and RPR. \n\n*** Ensure we have a copy of neuroimaging to review. \n\nIf not done, consider detailed neuropsychological evaluation.",
                'demRiskPlan' : {
                    'vascRisk' : "The need to attend to any vascular risk factors was discussed, as was the importance of quality sleep, a healthy diet, and mental and physical activity.",
                    'depAnx' :  "Treatment of any depression or anxiety symptoms, including speaking with a qualified and trained professional as well as consideration of an antidepressant, was discussed." if svar('behDepress') != "Not present" or svar('behAnx') != "Not present" else
                        ""},
                'safety' : {
                    'gunStored' : "Ensure that guns are well secured." if svar('gunSecure') != "Yes" and svar('gunHome') != "No" else
                        "",
                    'violentBeh' : "Management of violent behavior was discussed, including a need to emphasize personal safety and when to involve police if necessary. \nPharmacological management of violent behavior was considered. \nNonpharmacological management of inappropriate behavior was also discussed." if svar('behViolent') != "No" else
                        "",
                    'suic' : "Suicidal ideation was discussed.  A prevention plan was considered, and information on suicidal ideation provided." if svar('behActSI') != "No" or svar('behPassiveSI') != "No" else
                        "",
                    'driving' : "Concerns regarding driving were discussed.  A driving evaluation was recommended." if (svar('adlDrive') == "Yes" and svar('adlMVA') == "Yes") or (surveyscore('FAQ') > 6 and svar('adlDrive') == "Yes") else
                        "",
                    'homeSafety' : "Concerns regarding safety at home were discussed, as well as strategies to maximize both safety and independence." if svar('adlHomeSafe') != "No" else
                        "",
                    'fallRisk' : "The risks of falls were discussed.  PT/OT and home fall precautions were considered." if svar('unsteadyFall') != "No" else
                        "",
                    'choking' : "A swallow evaluation was considered." if svar('choking') != "No" else
                        "",
                    'hoarding' : "Hoarding behavior was discussed." if  svar('hoarding') != "No" else
                        ""},
                'planning' : {
                    'advanDir' : "We discussed the importance of discussing advanced directives and power of attorney before these are really needed.  The need for proactive planning, including financial, legal and care planning was discussed.",
                    'lifePlanning' : "End of life planning including screening, treatment and placement was discussed. We discussed alternative care options available in the community. We discussed the need to explore care options well before there was a crisis. \nThe anticipated need for greater family involvement and support was discussed, as well as the need to recognize family limitations.\nThe idea of palliative or supportive care consults were introduced. \nRecognition of grieving process for the family members was discussed.  We recommended they explore what kinds of financial assistance for which the patient might qualify." if surveyscore('FAQ') > 15 or compositescore(DSRS) >= 18 else
                        "",
                    'hospicePlanning' : "The idea of hospice was introduced. \nAutopsy was discussed. \nWe discussed bereavement services." if surveyscore('FAQ') > 28 or compositescore(DSRS) > 34 else
                        ""},
                'genHealth' : "\nThe importance of a healthy diet was emphasized, as was the need for regular physical exercise (if discussed with primary care), mental activity and social activity.",
                'caregiver' : {
                    'caregiverHealth' : "The need for all caregivers involved to also attend to their own health was emphasized.  We recommend each caregiver having at least one day a week to do something they enjoy on their own.  Caregiver groups were also discussed.  Support groups were recommended. \nThe need for emergency plans was discussed in the event of caregiver illness.",
                    'socialWork' : "Additional resources such as social work were offered given signs of caregiver burnout." if surveyscore('zarit') >= 9 else
                        ""},
                'fallRiskPlan' : "The risks of falls were discussed.  PT/OT and home fall precautions were considered." if svar('unsteadyFall') != "No" else
                    "",
                'research' : "The patient was informed about available research studies through the University of Colorado in case they are interested in participating."},

            'hpi' : {
                'title': 'History of Present Illness',
                'ptSum' : str(svar('ptId')) + " is a ***" + "-year-old " + svar('ptSex').lower() + " with " + str(nvar('probSum')) + conditional(svar('probList'), "none of the above") + " who is visiting the Rocky Mountain Memory Disorders Clinic at the University of Colorado, Denver for a comprehensive care plan.  ",
                'introOutput' : nvar('careSum') + nvar('probSpeed') + "Problems have" + nvar('rapidWorsen') + "rapidly worsened in the last two weeks.  " + nvar('sigFlux'),
                'sectionTitle' : 'Dedicated Comprehensive Care Assessment',
                'medRos' : "Review of systems as per MA note: ***",
                'assessmentTitle' : "Assessment of Functional Capacity and Staging of Dementia Severity",
                'katzTitle' : "Katz Index of Independence in Activities of Daily Living (bADL assessment)" if nvar('adlList') else "",
                'katzOutput' : nvar('adlList'),
                'katzTable' : self.buildsurvey_dict('ADL'),
                'faqTitle' : "Functional Activities Questionnaire Review Flowsheet",
                'faqOutput' : nvar('faqList'),
                'faqTable' : self.buildsurvey_dict('FAQ'),
                'dsrsOutput' : nvar('DSRS'),
                'behDescTitle' : "Assessment of Neuropsychiatric Symptoms",
                'behDesc1' : "The NPI-Q (Kaufer et al. J Neuropsychiatry Clin Neurosci 2000) was filled. " + nvar('npiqList'),
                'behDesc2' : "He/she has "+ nvar('labile') + "been more emotionally labile.  ",
                'behList' : "There is " + nvar('behHygieneChange') + "report of changes in personal hygiene.  ",
                'hoardingSurv' : nvar('hoardingSurv'),
                'behDesc3' : "He/she " + nvar('passSI') + "thoughts that he/she would be better off dead or wished he/she could go to sleep and not wake up.  " + nvar('actSI2') + " he/she " + nvar('actSI1') + "thoughts of killing him/herself.  " + "He/she "+ nvar('violent') + "threatening, violent or had serious thoughts of harming others. ",
                'npiqTitle' : "NPI-Q Review Flowsheet",
                'npiqTable' : self.buildsurvey_dict('npiq'),
                'phqTitle' : "PHQ-9 Depression Scale",
                'phqTable' : self.buildsurvey_dict('phq'),
                'SafetySubjectiveTitle' : "Safety Evaluation",
                'subjSafetyOutput' : "An expanded Alzheimer's Association Safety Checklist was filled: " + '\n' + questionsandresponses(['adlDrive', 'adlMVA', 'adlHomeSafe', 'wandering', 'takingMeds', 'liveAlone', 'unsteadyFall', 'finMismanage', 'choking', 'hoarding']) +
                    "\nThese responses reveal the following: He/she is " + nvar('drive') + "driving.  There have been " + nvar('mva') + "accidents. There are " + nvar('safeHome') + "safety concerns at home. " + nvar('wandering') + nvar('medications') + nvar('liveAlone') + nvar('fallRiskAssess' ) + nvar('finMismanage') + nvar('choking') + nvar('hoarding'),
                'careDescTitle' : "Caregiver Evaluation",
                'careDescOutput' : nvar('careDesc1') + nvar('careDesc2') + '\nNeeds: *** ' + nvar('needsAddressed') + '\nSocial supports:  *** \n' +
                    'Ability and willingness to continue taking on and sustain caregiving tasks:  *** '},

            'ACPlan' : {
                'pallEOLtitle' : "Palliative Care Needs and End-of-life Checklists",
                'wishes' : 'Have wishes or desires for end-of-life care been discussed?:  *** ',
                'attorneyFin' : "Is a power of attorney in place for financial needs?:  *** ",
                'attorneyDec' : "Is a power of attorney in place for health care decisions?:  *** ",
                'pallHospCare' : "Is palliative care or hospice care appropriate for the patient?:   *** ",
            },

            'ObjectiveTitle' : "Objective",
            'dataImagingTitle' : 'Data / Imaging Review',
            'questionairesTitle' : "Further Questionnaires and Scores",
            'scores' : "Functional Activities Questionnaire Total: " + str(surveyscore('FAQ')) + " (out of 30 possible, " + str(surveypercent('FAQ')) + " % of maximum possible score)" +
                "\nZarit Caregiver Burden: " + str(surveyscore('zarit')) +" (out of 48 possible, " + str(surveypercent('zarit')) + " % of maximum possible score)"
                "\nNPI-Q Severity Total: " + str(surveyscore('npiq')) + " (out of 36, " + str(surveypercent('npiq')) + " % of maximum possible score)"
                + '\nDementia Severity Rating Scale: ' + str(compositescore(DSRS))+ " (out of 54 possible, " + str(round((compositescore(DSRS) / 54) * 100,2)) + " % of maximum possible score)",
            # ~ 'testTable': {
                # ~ 'hi':'there',
                # ~ 'what':'about',
                # ~ 'me':'?',
            # ~ },
        }




if __name__ == "__main__":
    pass
    # STP: I used these lines to make the ptnote.yaml files
    # ~ import yaml
    # ~ print(yaml.dump(Ptnote.standardizedsurveys))
    # ~ print(yaml.dump(Survey99483.sections))
